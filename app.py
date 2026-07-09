import json
import re
from datetime import date
from io import BytesIO
from pathlib import Path

import google.generativeai as genai
import streamlit as st
from docx import Document


TEMPLATE_KEYS = [
    "title",
    "date",
    "location",
    "reporter",
    "attendees",
    "agenda",
    "content",
    "tasks",
    "references",
]

MODEL_CANDIDATES = [
    "gemini-2.5-flash",
    "gemini-2.0-flash",
    "gemini-flash-latest",
]


def to_text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, (list, tuple)):
        return "\n".join(f"- {to_text(item).strip()}" for item in value if to_text(item).strip())
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, indent=2)
    return str(value)


def to_string_list(value) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return []
        # 개행 또는 쉼표 기반으로 카드 목록 표시에 적합하게 분리
        if "\n" in text:
            parts = [line.strip(" -\t") for line in text.splitlines()]
        else:
            parts = [part.strip() for part in text.split(",")]
        return [p for p in parts if p]
    if isinstance(value, (list, tuple)):
        return [to_text(item).strip() for item in value if to_text(item).strip()]
    return [to_text(value).strip()] if to_text(value).strip() else []


def find_default_template() -> Path | None:
    docx_files = sorted(Path(".").glob("*.docx"))
    return docx_files[0] if docx_files else None


def configure_gemini() -> tuple[bool, str]:
    api_key = st.secrets.get("GEMINI_API_KEY", "")
    if not api_key:
        return False, "st.secrets에 GEMINI_API_KEY가 없습니다."
    genai.configure(api_key=api_key)
    return True, ""


def choose_available_model_name() -> str:
    try:
        available = {
            model.name.replace("models/", "")
            for model in genai.list_models()
            if "generateContent" in getattr(model, "supported_generation_methods", [])
        }
    except Exception:
        # 목록 조회가 실패해도 대표 안정 모델로 시도
        return "gemini-2.0-flash"

    for candidate in MODEL_CANDIDATES:
        if candidate in available:
            return candidate
    raise RuntimeError(
        "사용 가능한 Gemini generateContent 모델을 찾지 못했습니다. "
        "Google AI Studio에서 API 키/모델 권한을 확인해 주세요."
    )


def extract_json_block(text: str) -> dict:
    text = text.strip()
    fenced = re.search(r"```(?:json)?\s*(\{.*\})\s*```", text, re.DOTALL)
    if fenced:
        text = fenced.group(1)
    return json.loads(text)


def build_prompt(raw_memo: str, default_date: str, title_hint: str, location_hint: str) -> str:
    return f"""
아래는 실제 회의 메모(대화체 포함)입니다. 이를 "개조식 단문형 회의결과보고서"로 정리하세요.

반드시 지켜야 할 기준:
1) 문장은 짧고 명확한 단문 위주
2) 핵심 안건과 결정사항이 한눈에 보이도록 구성
3) 과장/추측 금지, 메모 근거 중심
4) 출력은 오직 JSON 한 개만 반환 (추가 설명 금지)

JSON 스키마:
{{
  "title": "회의 제목(없으면 합리적으로 생성)",
  "date": "YYYY-MM-DD 형식, 없으면 {default_date}",
  "location": "회의 장소/채널, 없으면 {location_hint or '미정'}",
  "reporter": "작성자(모르면 미기재)",
  "attendees": "참석자 목록 (쉼표 구분)",
  "agenda": "주요안건을 개조식으로, 각 항목은 '- '로 시작",
  "content": "회의 내용 요약 (개조식 단문형, 핵심 근거 포함)",
  "tasks": "결정사항/후속조치 (담당/기한 있으면 포함, 각 항목 '- ')",
  "references": "참고자료/링크/첨부 목록 (없으면 '없음')",
  "major_agendas": ["시선이 가는 짧은 안건 3~6개"],
  "decisions": ["핵심 결정사항 3~8개"]
}}

제목 힌트: {title_hint or '없음'}
장소 힌트: {location_hint or '없음'}

[회의 메모 원문]
{raw_memo}
""".strip()


def summarize_with_gemini(raw_memo: str, title_hint: str, location_hint: str) -> dict:
    model_name = choose_available_model_name()
    model = genai.GenerativeModel(model_name)
    default_date = date.today().isoformat()
    prompt = build_prompt(raw_memo, default_date, title_hint, location_hint)
    response = model.generate_content(prompt)
    data = extract_json_block(response.text)

    # 템플릿 필드 기본값 보정
    for key in TEMPLATE_KEYS:
        data.setdefault(key, "")
    data.setdefault("major_agendas", [])
    data.setdefault("decisions", [])
    data["used_model"] = model_name
    return data


def replace_text_preserving_style(text: str, mapping: dict[str, str]) -> str:
    for key, value in mapping.items():
        text = text.replace(f"{{{{{key}}}}}", value)
    return text


def replace_placeholders_in_doc(document: Document, mapping: dict[str, str]) -> None:
    def replace_in_paragraph(paragraph) -> None:
        text_runs = [run for run in paragraph.runs if run.text]
        if not text_runs:
            return

        original = "".join(run.text for run in text_runs)
        replaced = replace_text_preserving_style(original, mapping)
        if replaced == original:
            return

        # paragraph.text를 직접 재할당하면 상단 템플릿(로고/도형)까지 손상될 수 있어
        # 텍스트 run만 최소 수정한다.
        text_runs[0].text = replaced
        for run in text_runs[1:]:
            run.text = ""

    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph)


def clear_unresolved_placeholders(document: Document) -> list[str]:
    unresolved: set[str] = set()
    pattern = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")

    def clear_in_paragraph(paragraph) -> None:
        text_runs = [run for run in paragraph.runs if run.text]
        if not text_runs:
            return

        text = "".join(run.text for run in text_runs)
        hits = pattern.findall(text)
        if hits:
            unresolved.update(hit.strip() for hit in hits)
            cleaned = pattern.sub("", text)
            text_runs[0].text = cleaned
            for run in text_runs[1:]:
                run.text = ""

    for paragraph in document.paragraphs:
        clear_in_paragraph(paragraph)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    clear_in_paragraph(paragraph)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            clear_in_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            clear_in_paragraph(paragraph)

    return sorted(unresolved)


def build_output_filename(fields: dict[str, str], output_dir: Path) -> str:
    raw_date = (fields.get("date") or "").strip()
    date_part = re.sub(r"[^0-9]", "", raw_date)[:8] or date.today().strftime("%Y%m%d")

    raw_title = (fields.get("title") or "").strip() or "회의결과보고서"
    title_part = re.sub(r"[\\\\/:*?\"<>|]", "_", raw_title)
    title_part = re.sub(r"\s+", "_", title_part).strip("_") or "회의결과보고서"

    base = f"{date_part}_{title_part}"
    pattern = re.compile(rf"^{re.escape(base)}_v(\d+)\.docx$", re.IGNORECASE)

    max_version = 0
    if output_dir.exists():
        for file_path in output_dir.glob(f"{base}_v*.docx"):
            match = pattern.match(file_path.name)
            if match:
                max_version = max(max_version, int(match.group(1)))

    next_version = max_version + 1
    return f"{base}_v{next_version:02d}.docx"


def render_editor_fields() -> dict[str, str]:
    st.subheader("AI 정리 결과 (수정 가능)")
    fields: dict[str, str] = {}

    fields["title"] = st.text_input("제목", key="f_title")
    fields["date"] = st.text_input("일자", key="f_date")
    fields["location"] = st.text_input("장소/채널", key="f_location")
    fields["reporter"] = st.text_input("작성자", key="f_reporter")
    fields["attendees"] = st.text_area("참석자", height=80, key="f_attendees")
    fields["agenda"] = st.text_area("주요안건", height=130, key="f_agenda")
    fields["content"] = st.text_area("회의 내용 요약", height=220, key="f_content")
    fields["tasks"] = st.text_area("결정사항/후속조치", height=180, key="f_tasks")
    fields["references"] = st.text_area("참고자료", height=80, key="f_references")
    return fields


def load_fields_into_session(data: dict) -> None:
    st.session_state["f_title"] = to_text(data.get("title", ""))
    st.session_state["f_date"] = to_text(data.get("date", ""))
    st.session_state["f_location"] = to_text(data.get("location", ""))
    st.session_state["f_reporter"] = to_text(data.get("reporter", ""))
    st.session_state["f_attendees"] = to_text(data.get("attendees", ""))
    st.session_state["f_agenda"] = to_text(data.get("agenda", ""))
    st.session_state["f_content"] = to_text(data.get("content", ""))
    st.session_state["f_tasks"] = to_text(data.get("tasks", ""))
    st.session_state["f_references"] = to_text(data.get("references", ""))
    st.session_state["major_agendas"] = to_string_list(data.get("major_agendas", []))
    st.session_state["decisions"] = to_string_list(data.get("decisions", []))
    st.session_state["used_model"] = to_text(data.get("used_model", ""))


def main() -> None:
    st.set_page_config(page_title="회의결과보고서 자동생성기", page_icon="📝", layout="wide")
    st.title("📝 회의결과보고서 자동생성기")
    st.caption("회의 메모를 입력하면 Gemini가 개조식 회의결과보고서로 정리하고 Word 템플릿을 채웁니다.")

    configured, err = configure_gemini()
    if not configured:
        st.error(err)
        st.info("`.streamlit/secrets.toml` 또는 Streamlit Cloud secrets에 API 키를 설정하세요.")
        st.stop()

    default_template = find_default_template()
    st.markdown("### 1) 회의 메모 입력")
    c1, c2 = st.columns([2, 1])
    with c1:
        raw_memo = st.text_area(
            "회의 메모(대화내용 그대로)",
            height=260,
            placeholder="예) 김과장: 3분기 예산은 마케팅 10% 증액...\n이대리: 개발 일정은 8월 2주차 완료 목표...",
        )
    with c2:
        title_hint = st.text_input("회의 제목 힌트(선택)")
        location_hint = st.text_input("장소/회의채널 힌트(선택)")
        st.write("")
        if default_template:
            st.success(f"기본 템플릿 인식: `{default_template.name}`")
        else:
            st.warning("작업 폴더에 기본 .docx 템플릿이 없습니다.")

    st.markdown("### 2) 템플릿 선택")
    uploaded = st.file_uploader("템플릿(.docx) 업로드 (선택)", type=["docx"])

    if st.button("AI로 회의결과 정리", type="primary", use_container_width=True):
        if not raw_memo.strip():
            st.warning("회의 메모를 먼저 입력하세요.")
            st.stop()
        with st.spinner("Gemini가 회의 메모를 정리하는 중입니다..."):
            try:
                data = summarize_with_gemini(raw_memo, title_hint, location_hint)
                load_fields_into_session(data)
                st.success("AI 정리가 완료되었습니다. 아래에서 수정 후 Word 파일을 생성하세요.")
            except Exception as e:
                st.error(f"AI 정리 중 오류가 발생했습니다: {e}")

    if "f_title" in st.session_state:
        used_model = st.session_state.get("used_model", "")
        if used_model:
            st.caption(f"사용 모델: `{used_model}`")
        st.markdown("### 3) 주요안건/결정사항 한눈에 보기")
        a_col, d_col = st.columns(2)
        with a_col:
            st.info("주요안건")
            agendas = st.session_state.get("major_agendas", [])
            if agendas:
                for item in agendas:
                    st.markdown(f"- **{item}**")
            else:
                st.markdown("- 없음")
        with d_col:
            st.success("결정사항")
            decisions = st.session_state.get("decisions", [])
            if decisions:
                for item in decisions:
                    st.markdown(f"- **{item}**")
            else:
                st.markdown("- 없음")

        st.markdown("### 4) 결과 수정 및 Word 생성")
        fields = render_editor_fields()

        selected_template_bytes = None
        selected_template_name = None
        if uploaded is not None:
            selected_template_bytes = uploaded.getvalue()
            selected_template_name = uploaded.name
        elif default_template is not None:
            selected_template_bytes = default_template.read_bytes()
            selected_template_name = default_template.name

        if selected_template_bytes is None:
            st.warning("템플릿 파일이 필요합니다. .docx 파일을 업로드하세요.")
            st.stop()

        if st.button("Word 회의결과보고서 생성", use_container_width=True):
            try:
                doc = Document(BytesIO(selected_template_bytes))
                replace_placeholders_in_doc(doc, fields)
                unresolved = clear_unresolved_placeholders(doc)
                out = BytesIO()
                doc.save(out)
                out.seek(0)

                output_dir = Path("generated_reports")
                output_dir.mkdir(exist_ok=True)
                output_name = build_output_filename(fields, output_dir)
                output_path = output_dir / output_name
                output_bytes = out.getvalue()
                output_path.write_bytes(output_bytes)

                st.session_state["generated_doc_bytes"] = output_bytes
                st.session_state["generated_doc_name"] = output_name
                st.session_state["generated_doc_path"] = str(output_path)

                st.success(f"템플릿 `{selected_template_name}` 기반 문서 생성 완료: `{output_path}`")
                if unresolved:
                    st.warning(
                        "템플릿에 남아있던 미매핑 플레이스홀더를 빈값으로 정리했습니다: "
                        + ", ".join(unresolved)
                    )
            except Exception as e:
                st.error(f"Word 생성 중 오류가 발생했습니다: {e}")

        if "generated_doc_bytes" in st.session_state:
            st.download_button(
                "📥 Word 파일 다운로드",
                data=st.session_state["generated_doc_bytes"],
                file_name=st.session_state.get("generated_doc_name", "회의결과보고서_자동생성.docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )


if __name__ == "__main__":
    main()
